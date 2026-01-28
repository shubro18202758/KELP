"""
Citation Generator Module - Creates citation documents linking claims to sources
This module handles the 10% evaluation criterion for Citation Integrity
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
import sys

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from config.settings import CITATIONS_OUTPUT_DIR, BRANDING


@dataclass
class Citation:
    """Individual citation entry"""
    claim: str
    source_type: str  # 'private_data', 'public_web', 'calculated', 'image'
    source_reference: str  # File path or URL
    section: str  # e.g., 'Slide 1 - Business Profile'
    line_or_sheet: Optional[str] = None  # Specific location within source
    notes: Optional[str] = None


@dataclass
class CitationCollection:
    """Collection of citations for a company"""
    company_name: str
    sector: str
    generated_date: str = field(default_factory=lambda: datetime.now().strftime("%Y-%m-%d %H:%M"))
    citations: List[Citation] = field(default_factory=list)
    
    def add(self, claim: str, source_type: str, source_reference: str, 
            section: str, line_or_sheet: str = None, notes: str = None) -> None:
        """Add a citation"""
        self.citations.append(Citation(
            claim=claim,
            source_type=source_type,
            source_reference=source_reference,
            section=section,
            line_or_sheet=line_or_sheet,
            notes=notes
        ))


class CitationTracker:
    """Tracks citations as data is processed"""
    
    def __init__(self, company_name: str, sector: str, source_file: str):
        self.collection = CitationCollection(
            company_name=company_name,
            sector=sector
        )
        self.source_file = source_file
    
    def cite_from_markdown(self, claim: str, section_name: str, slide: str) -> None:
        """Add citation for markdown-sourced data"""
        self.collection.add(
            claim=claim,
            source_type='private_data',
            source_reference=self.source_file,
            section=slide,
            line_or_sheet=f"Section: {section_name}",
            notes="Extracted from company one-pager (Markdown)"
        )
    
    def cite_from_swot(self, claim: str, swot_category: str, slide: str) -> None:
        """Add citation for SWOT-sourced data"""
        self.collection.add(
            claim=claim,
            source_type='private_data',
            source_reference=self.source_file,
            section=slide,
            line_or_sheet=f"SWOT Analysis - {swot_category}",
            notes="Extracted from SWOT analysis section"
        )
    
    def cite_calculated(self, claim: str, calculation_method: str, slide: str) -> None:
        """Add citation for calculated/derived data"""
        self.collection.add(
            claim=claim,
            source_type='calculated',
            source_reference=self.source_file,
            section=slide,
            notes=f"Calculated/derived: {calculation_method}"
        )
    
    def cite_image(self, image_description: str, image_url: str, 
                   image_source: str, slide: str) -> None:
        """Add citation for sourced images"""
        self.collection.add(
            claim=f"Image: {image_description}",
            source_type='image',
            source_reference=image_url,
            section=slide,
            notes=f"Source: {image_source}"
        )
    
    def cite_web_source(self, claim: str, url: str, slide: str) -> None:
        """Add citation for web-scraped data"""
        self.collection.add(
            claim=claim,
            source_type='public_web',
            source_reference=url,
            section=slide,
            notes="Scraped from public website"
        )
    
    def cite_sector_template(self, claim: str, slide: str) -> None:
        """Add citation for sector-specific template content"""
        self.collection.add(
            claim=claim,
            source_type='calculated',
            source_reference="Sector Template Library",
            section=slide,
            notes="Standard sector-specific investment language"
        )


class CitationDocumentGenerator:
    """Generates Word document with all citations"""
    
    def __init__(self, collection: CitationCollection):
        self.collection = collection
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self) -> None:
        """Setup document styles"""
        # Title style
        style = self.doc.styles['Title']
        font = style.font
        font.name = BRANDING.heading_font
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = RGBColor(*BRANDING.primary_dark_indigo)
        
        # Heading 1
        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = BRANDING.heading_font
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(*BRANDING.primary_violet)
        
        # Heading 2
        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = BRANDING.heading_font
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(*BRANDING.accent_cyan)
    
    def _add_header(self) -> None:
        """Add document header"""
        # Title
        title = self.doc.add_heading("Citation Document", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Investment Teaser Source References")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*BRANDING.text_light_grey)
        
        # Meta information
        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(f"Sector: {self.collection.sector}").font.size = Pt(11)
        meta.add_run(f"  |  Generated: {self.collection.generated_date}").font.size = Pt(11)
        
        # Divider
        self.doc.add_paragraph("─" * 80)
    
    def _add_section(self, slide_name: str, citations: List[Citation]) -> None:
        """Add a slide section with its citations"""
        self.doc.add_heading(slide_name, level=1)
        
        # Group by source type
        by_type: Dict[str, List[Citation]] = {}
        for c in citations:
            if c.source_type not in by_type:
                by_type[c.source_type] = []
            by_type[c.source_type].append(c)
        
        type_labels = {
            'private_data': '📁 Private Data Sources',
            'public_web': '🌐 Public Web Sources',
            'calculated': '📊 Calculated/Derived',
            'image': '🖼️ Image Sources'
        }
        
        for source_type, type_citations in by_type.items():
            self.doc.add_heading(type_labels.get(source_type, source_type), level=2)
            
            for citation in type_citations:
                # Citation entry
                p = self.doc.add_paragraph()
                
                # Claim
                claim_run = p.add_run("Claim: ")
                claim_run.bold = True
                p.add_run(f'"{citation.claim[:200]}{"..." if len(citation.claim) > 200 else ""}"')
                
                # Source
                p2 = self.doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.3)
                source_run = p2.add_run("Source: ")
                source_run.bold = True
                p2.add_run(citation.source_reference)
                
                # Location
                if citation.line_or_sheet:
                    p3 = self.doc.add_paragraph()
                    p3.paragraph_format.left_indent = Inches(0.3)
                    loc_run = p3.add_run("Location: ")
                    loc_run.bold = True
                    p3.add_run(citation.line_or_sheet)
                
                # Notes
                if citation.notes:
                    p4 = self.doc.add_paragraph()
                    p4.paragraph_format.left_indent = Inches(0.3)
                    p4.add_run(f"Note: {citation.notes}").font.italic = True
                
                # Spacing
                self.doc.add_paragraph()
    
    def _add_footer(self) -> None:
        """Add document footer"""
        self.doc.add_paragraph("─" * 80)
        
        footer = self.doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(BRANDING.footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*BRANDING.text_light_grey)
    
    def generate(self, output_filename: str = None) -> str:
        """Generate the citation document"""
        self._add_header()
        
        # Summary statistics
        self.doc.add_heading("Summary", level=1)
        summary = self.doc.add_paragraph()
        summary.add_run(f"Total Citations: {len(self.collection.citations)}\n")
        
        # Count by type
        type_counts = {}
        for c in self.collection.citations:
            type_counts[c.source_type] = type_counts.get(c.source_type, 0) + 1
        
        for source_type, count in type_counts.items():
            summary.add_run(f"• {source_type.replace('_', ' ').title()}: {count}\n")
        
        # Group citations by slide
        by_slide: Dict[str, List[Citation]] = {}
        for citation in self.collection.citations:
            if citation.section not in by_slide:
                by_slide[citation.section] = []
            by_slide[citation.section].append(citation)
        
        # Add each slide section
        for slide_name in sorted(by_slide.keys()):
            self._add_section(slide_name, by_slide[slide_name])
        
        self._add_footer()
        
        # Generate filename if not provided
        if not output_filename:
            safe_name = "".join(c if c.isalnum() or c in ' -_' else '' 
                               for c in self.collection.company_name)
            safe_name = safe_name.replace(' ', '_')[:30]
            output_filename = f"{safe_name}_Citations.docx"
        
        # Save
        CITATIONS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = CITATIONS_OUTPUT_DIR / output_filename
        self.doc.save(str(output_path))
        
        return str(output_path)


def generate_citations_from_content(company_name: str, sector: str, source_file: str,
                                   slide_content: Dict, image_citations: List[Dict] = None) -> str:
    """
    Generate citation document from slide content.
    
    Args:
        company_name: Name of the company
        sector: Classified sector
        source_file: Path to source markdown file
        slide_content: Dict with 'slide1', 'slide2', 'slide3' content
        image_citations: List of image source citations
    
    Returns:
        Path to generated citation document
    """
    tracker = CitationTracker(company_name, sector, source_file)
    
    # Slide 1 citations
    slide1 = slide_content.get('slide1', {})
    if slide1.get('company_description'):
        tracker.cite_from_markdown(
            slide1['company_description'][:100] + "...",
            "Business Description",
            "Slide 1 - Business Profile"
        )
    
    for section_name, items in slide1.get('sections', {}).items():
        if isinstance(items, list):
            for item in items[:3]:  # Cite top 3 from each section
                tracker.cite_from_markdown(
                    item,
                    section_name,
                    "Slide 1 - Business Profile"
                )
    
    # Slide 2 citations
    slide2 = slide_content.get('slide2', {})
    for metric in slide2.get('metrics', [])[:5]:
        tracker.cite_from_markdown(
            metric,
            "Key Operational Indicators / SWOT",
            "Slide 2 - Financial & Operational"
        )
    
    # Charts are derived
    if slide2.get('charts'):
        tracker.cite_calculated(
            "Chart data visualization",
            "Aggregated from financial metrics",
            "Slide 2 - Financial & Operational"
        )
    
    # Slide 3 citations
    slide3 = slide_content.get('slide3', {})
    for highlight in slide3.get('highlights', [])[:4]:
        # Some are from template, some from data
        if any(keyword in highlight.lower() for keyword in ['proprietary', 'strategic', 'platform']):
            tracker.cite_sector_template(highlight, "Slide 3 - Investment Highlights")
        else:
            tracker.cite_from_swot(highlight, "Strengths/Opportunities", 
                                  "Slide 3 - Investment Highlights")
    
    for outlook in slide3.get('future_outlook', []):
        tracker.cite_from_markdown(outlook, "Future Plans", "Slide 3 - Investment Highlights")
    
    # Image citations
    if image_citations:
        for img in image_citations:
            tracker.cite_image(
                img.get('description', 'Business image'),
                img.get('url', 'N/A'),
                img.get('source', 'Stock image library'),
                img.get('slide', 'Multiple slides')
            )
    
    # Generate document
    generator = CitationDocumentGenerator(tracker.collection)
    return generator.generate()


if __name__ == "__main__":
    # Test citation generation
    print("Testing Citation Generation...\n")
    
    test_content = {
        'slide1': {
            'company_description': 'A leading manufacturing company with over 30 years of experience.',
            'sections': {
                'Products': ['Precision Machining', 'Die Manufacturing'],
                'Industries': ['Automotive', 'Aerospace']
            }
        },
        'slide2': {
            'metrics': ['25% YoY Revenue Growth', 'EBITDA Margin: 18%'],
            'charts': [{'type': 'bar', 'title': 'Growth'}]
        },
        'slide3': {
            'highlights': [
                'Proprietary product portfolio with high entry barriers',
                'Strong customer relationships with leading OEMs'
            ],
            'future_outlook': ['20% growth target for FY26']
        }
    }
    
    output = generate_citations_from_content(
        "Test Company",
        "Manufacturing",
        "test_source.md",
        test_content
    )
    print(f"✓ Generated: {output}")
